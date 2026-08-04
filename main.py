from typing import List, Optional
from fastapi import FastAPI, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import psycopg2
from psycopg2.extras import RealDictCursor
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

app = FastAPI(title="API Suivi Budgétaire Ministériel")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En production, mettez l'URL de votre app Firebase
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DB_CONFIG = {
    "dbname": "budget_db",
    "user": "postgres",
    "password": "zougalB613",
    "host": "localhost",
    "port": "5432"
}

def get_db_connection():
    try:
        return psycopg2.connect(**DB_CONFIG, cursor_factory=RealDictCursor)
    except Exception:
        return None

# ============================================================
# MODÈLES PYDANTIC
# ============================================================
class LoginRequest(BaseModel):
    identifiant: str  # P01, P02, P03, P04, ADMIN
    mot_de_passe: str
    exercice: int = 2026

class LigneBudgetaire(BaseModel):
    id: int
    label: str
    lfi: float
    ajustement_lfr: float = 0
    engagements: float = 0
    paiements: float = 0

class ExecutionIcp(BaseModel):
    id: int
    realise: float

class ExecutionVentilation(BaseModel):
    nature: str
    dotation_lfi: float
    ajustement_lfr: float = 0
    engagements: float = 0
    paiements: float = 0

class SaisieTrimestriellePayload(BaseModel):
    programme_id: int
    exercice: int = 2026
    trimestre: str
    lignes: List[LigneBudgetaire]
    icps: List[ExecutionIcp]
    ventilation: Optional[List[ExecutionVentilation]] = []

class LfrTogglePayload(BaseModel):
    exercice: int
    lfr_active: bool

# Métadonnées en mémoire si pas de BD physique
LFR_STATUS_BY_EXERCICE = { 2024: False, 2025: False, 2026: False, 2027: False }
CIBLES_TRIMESTRE = {"T1": 25, "T2": 50, "T3": 75, "T4": 100}

SEUIL_ALERTE = 5
SEUIL_CRITIQUE = 15

# Données de secours par défaut si la connexion BD est absente
HISTORIQUE_EXERCICES = {
    2026: [
        {"id": 1, "code": "P01", "nom": "Enseignement Supérieur", "t1": 10.0, "t2": 21.0, "t3": 0, "t4": 0, "lignes": [{"id": 101, "label": "Construction et équipement d'écoles", "lfi": 1000000000, "ajustement_lfr": 150000000, "engagements": 300000000, "paiements": 210000000}], "indicateurs": [{"id": 201, "nom": "Taux de scolarisation", "unite": "%", "cible": 95, "realise": 88, "inverse": False}], "ventilation": [{"nature": "Dépenses de personnel", "dotation_lfi": 800000000, "ajustement_lfr": 50000000, "engagements": 300000000, "paiements": 210000000}]},
        {"id": 2, "code": "P02", "nom": "Santé Publique", "t1": 8.5, "t2": 18.5, "t3": 0, "t4": 0, "lignes": [{"id": 103, "label": "Approvisionnement en médicaments", "lfi": 800000000, "ajustement_lfr": -50000000, "engagements": 250000000, "paiements": 148000000}], "indicateurs": [{"id": 202, "nom": "Taux de couverture vaccinale", "unite": "%", "cible": 90, "realise": 82, "inverse": False}], "ventilation": [{"nature": "Transferts courants", "dotation_lfi": 600000000, "ajustement_lfr": 0, "engagements": 150000000, "paiements": 111000000}]},
        {"id": 3, "code": "P03", "nom": "Infrastructures Routières", "t1": 12.0, "t2": 30.0, "t3": 0, "t4": 0, "lignes": [{"id": 105, "label": "Entretien du réseau routier", "lfi": 2000000000, "ajustement_lfr": 200000000, "engagements": 800000000, "paiements": 600000000}], "indicateurs": [{"id": 203, "nom": "Routes bitumées ou entretenues", "unite": "Km", "cible": 150, "realise": 45, "inverse": False}], "ventilation": [{"nature": "Investissements exécutés par l'État", "dotation_lfi": 3000000000, "ajustement_lfr": 200000000, "engagements": 1200000000, "paiements": 900000000}]},
        {"id": 4, "code": "P04", "nom": "Gouvernance & Administration", "t1": 9.0, "t2": 19.7, "t3": 0, "t4": 0, "lignes": [{"id": 107, "label": "Modernisation des services", "lfi": 600000000, "ajustement_lfr": 0, "engagements": 180000000, "paiements": 118200000}], "indicateurs": [{"id": 204, "nom": "Taux de dématérialisation", "unite": "%", "cible": 80, "realise": 65, "inverse": False}], "ventilation": [{"nature": "Dépenses de fonctionnement", "dotation_lfi": 1000000000, "ajustement_lfr": 0, "engagements": 280000000, "paiements": 197000000}]}
    ],
    2027: [
        {"id": 1, "code": "P01", "nom": "Éducation & Formation Pro", "t1": 0, "t2": 0, "t3": 0, "t4": 0, "lignes": [], "indicateurs": [], "ventilation": []},
        {"id": 2, "code": "P02", "nom": "Santé & Protection Sociale", "t1": 0, "t2": 0, "t3": 0, "t4": 0, "lignes": [], "indicateurs": [], "ventilation": []},
        {"id": 3, "code": "P03", "nom": "Transports & Désenclavement", "t1": 0, "t2": 0, "t3": 0, "t4": 0, "lignes": [], "indicateurs": [], "ventilation": []},
        {"id": 4, "code": "P04", "nom": "Pilotage & Support Ministériel", "t1": 0, "t2": 0, "t3": 0, "t4": 0, "lignes": [], "indicateurs": [], "ventilation": []}
    ]
}

# ============================================================
# ENDPOINT D'AUTHENTIFICATION SIMPLE (P01 - P04)
# ============================================================
@app.post("/api/login")
def login(payload: LoginRequest):
    ident = payload.identifiant.upper().strip()
    pwd = payload.mot_de_passe.strip()
    
    conn = get_db_connection()
    if conn:
        try:
            cur = conn.cursor()
            cur.execute("SELECT identifiant, role FROM utilisateurs WHERE identifiant = %s AND mot_de_passe = %s;", (ident, pwd))
            user = cur.fetchone()
            if not user:
                cur.close(); conn.close()
                raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Identifiant ou mot de passe incorrect")
            
            role = user['role']
            nom_programme = "Tous les programmes"
            
            if role == 'RESPONSABLE':
                cur.execute("SELECT nom FROM programmes WHERE code_programme = %s AND exercice = %s;", (ident, payload.exercice))
                prog = cur.fetchone()
                if prog:
                    nom_programme = prog['nom']
            
            cur.close(); conn.close()
            return {"status": "success", "identifiant": ident, "role": role, "exercice": payload.exercice, "nom_programme": nom_programme}
        except Exception:
            if conn: conn.close()

    # Fallback si BD non connectée
    if pwd == "pass2026" or pwd == "admin2026":
        role = "ADMIN" if ident == "ADMIN" else "RESPONSABLE"
        progs_ex = HISTORIQUE_EXERCICES.get(payload.exercice, HISTORIQUE_EXERCICES[2026])
        prog_match = next((p for p in progs_ex if p["code"] == ident), None)
        nom_prog = prog_match["nom"] if prog_match else "Tous les programmes"
        return {"status": "success", "identifiant": ident, "role": role, "exercice": payload.exercice, "nom_programme": nom_prog}

    raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Identifiant ou mot de passe incorrect")


# ============================================================
# LECTURE ET GESTION DES PROGRAMMES
# ============================================================
def _fetch_programmes_data(exercice: int = 2026, trimestre: str = "T2"):
    conn = get_db_connection()
    lfr_status = LFR_STATUS_BY_EXERCICE.get(exercice, False)

    if not conn:
        data = HISTORIQUE_EXERCICES.get(exercice, HISTORIQUE_EXERCICES[2026])
        for p in data: p['lfr_active'] = lfr_status
        return data

    try:
        cur = conn.cursor()
        cur.execute("SELECT lfr_active FROM config_exercice WHERE exercice = %s;", (exercice,))
        res = cur.fetchone()
        if res:
            lfr_status = res['lfr_active']
            LFR_STATUS_BY_EXERCICE[exercice] = lfr_status

        cur.execute("SELECT id, code_programme AS code, nom FROM programmes WHERE exercice = %s ORDER BY code_programme;", (exercice,))
        programmes = cur.fetchall()

        for p in programmes:
            p['lfr_active'] = lfr_status
            cur.execute("SELECT id, label, COALESCE(lfi, ouverts) AS lfi, COALESCE(ajustement_lfr, 0) AS ajustement_lfr, COALESCE(engagements, 0) AS engagements, COALESCE(paiements, 0) AS paiements FROM lignes_budgetaires WHERE programme_id = %s ORDER BY id;", (p['id'],))
            p['lignes'] = cur.fetchall()

            cur.execute("SELECT id, nom, unite, cible_annuelle AS cible, COALESCE(realise, 0) AS realise, inverse FROM indicateurs WHERE programme_id = %s ORDER BY id;", (p['id'],))
            p['indicateurs'] = cur.fetchall()

            cur.execute("SELECT nature_economique AS nature, COALESCE(dotation_lfi, dotation) AS dotation_lfi, COALESCE(ajustement_lfr, 0) AS ajustement_lfr, COALESCE(engagements, 0) AS engagements, COALESCE(paiements, 0) AS paiements FROM ventilation_economique WHERE programme_id = %s;", (p['id'],))
            p['ventilation'] = cur.fetchall()

        cur.close(); conn.close()
        return programmes
    except Exception:
        if conn: conn.close()
        return HISTORIQUE_EXERCICES.get(exercice, [])

@app.get("/api/programmes")
def get_programmes(exercice: int = 2026, trimestre: str = "T2"):
    return _fetch_programmes_data(exercice, trimestre)

# ============================================================
# CARTOGRAPHIE DES RISQUES D'EXÉCUTION
# ============================================================
def _taux_et_assiette(ligne: dict, lfr_active: bool):
    lfi = float(ligne.get("lfi", 0) or 0)
    lfr = float(ligne.get("ajustement_lfr", 0) or 0) if lfr_active else 0.0
    assiette = lfi + lfr
    paiements = float(ligne.get("paiements", 0) or 0)
    taux = (paiements / assiette * 100) if assiette > 0 else 0.0
    return taux, assiette

def _niveau_risque(taux: float, cible: float):
    ecart = taux - cible
    if ecart <= -SEUIL_CRITIQUE: return "Critique"
    if ecart <= -SEUIL_ALERTE: return "Alerte"
    return "Normal"

def _calculer_risques(exercice: int, trimestre: str):
    cible = CIBLES_TRIMESTRE.get(trimestre, 50)
    programmes = _fetch_programmes_data(exercice, trimestre)
    lfr_active = LFR_STATUS_BY_EXERCICE.get(exercice, False)

    lignes_risque = []
    for p in programmes:
        nom_prog = p.get("nom", "Programme")
        for l in (p.get("lignes") or []):
            taux, assiette = _taux_et_assiette(l, lfr_active)
            if assiette <= 0: continue
            niveau = _niveau_risque(taux, cible)
            if niveau == "Normal": continue
            montant_a_risque = assiette - float(l.get("paiements", 0) or 0)
            lignes_risque.append({
                "programme": nom_prog,
                "ligne_id": l.get("id"),
                "label": l.get("label"),
                "montant_prevu": round(assiette, 2),
                "montant_execute": round(float(l.get("paiements", 0) or 0), 2),
                "montant_a_risque": round(max(montant_a_risque, 0), 2),
                "taux_execution": round(taux, 1),
                "cible": cible,
                "ecart": round(taux - cible, 1),
                "niveau_risque": niveau,
            })

    ordre = {"Critique": 0, "Alerte": 1}
    lignes_risque.sort(key=lambda x: (ordre.get(x["niveau_risque"], 2), -x["montant_a_risque"]))
    return lignes_risque

@app.get("/api/risques")
def get_risques(exercice: int = 2026, trimestre: str = "T2"):
    lignes = _calculer_risques(exercice, trimestre)
    return {
        "exercice": exercice,
        "trimestre": trimestre,
        "cible_trimestre": CIBLES_TRIMESTRE.get(trimestre, 50),
        "nb_critique": sum(1 for l in lignes if l["niveau_risque"] == "Critique"),
        "nb_alerte": sum(1 for l in lignes if l["niveau_risque"] == "Alerte"),
        "montant_total_a_risque": round(sum(l["montant_a_risque"] for l in lignes), 2),
        "lignes": lignes,
    }

# ============================================================
# STYLE & FORMATAGE DES TABLEAUX DOCX (ARIAL & JUSTIFIÉ)
# ============================================================
HEX_HEADER_BLUE = "1F4E79"
HEX_SUBHEADER = "EFF6FF"
HEX_BORDER = "D9D9D9"

def _set_cell_background(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def _set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def _appliquer_style_tableau_image(table):
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{HEX_BORDER}"/>
        </w:tblBorders>
    '''))

def _format_cellule_donnees(cell, texte: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    _set_cell_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(texte)
    r.bold = bold
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)

def _ajouter_source_tableau(doc, source_text: str = "SEB Ministériel, Direction des Affaires Financières"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Source : {source_text}")
    r.font.name = 'Arial'
    r.font.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

def _ajouter_analyse_paragraphe(doc, titre: str, texte_analyse: str):
    h = doc.add_heading(f"Analyse qualitative : {titre}", level=3)
    for run in h.runs: run.font.name = 'Arial'
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(texte_analyse)
    r.font.name = 'Arial'
    r.font.size = Pt(10)

def _fmt_fcfa(n: float) -> str: return f"{round(n or 0):,}".replace(",", " ") + " FCFA"
def _fmt_pct(n: float) -> str: return f"{n:.1f}%".replace(".", ",")

# ============================================================
# GENERATION DE RAPPORTS WORD (.DOCX)
# ============================================================
def _generer_rapport_docx(exercice: int, trimestre: str) -> io.BytesIO:
    cible = CIBLES_TRIMESTRE.get(trimestre, 50)
    programmes = _fetch_programmes_data(exercice, trimestre)
    lfr_active = LFR_STATUS_BY_EXERCICE.get(exercice, False)
    risques = _calculer_risques(exercice, trimestre)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8); section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    tot_lfi_global = sum(sum(float(l.get("lfi", 0) or 0) for l in p.get("lignes", [])) for p in programmes)

    def _add_heading_arial(text, level):
        h = doc.add_heading(text, level=level)
        for r in h.runs: r.font.name = 'Arial'
        return h

    # PAGE DE GARDE
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.add_run("RÉPUBLIQUE DU SÉNÉGAL\n").bold = True
    p_hdr.runs[0].font.name = 'Arial'; p_hdr.runs[0].font.size = Pt(13); p_hdr.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p_hdr.add_run("MINISTÈRE DE L'INTÉGRATION AFRICAINE ET DES AFFAIRES ÉTRANGÈRES\n").bold = True
    p_hdr.runs[1].font.name = 'Arial'; p_hdr.runs[1].font.size = Pt(11)

    tbl_garde = doc.add_table(rows=1, cols=1)
    tbl_garde.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_g = tbl_garde.cell(0, 0)
    _set_cell_background(cell_g, HEX_HEADER_BLUE)
    _set_cell_margins(cell_g, top=300, bottom=300, left=200, right=200)
    p_g = cell_g.paragraphs[0]; p_g.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rg1 = p_g.add_run("RAPPORT TRIMESTRIEL DE PERFORMANCE BUDGÉTAIRE\n")
    rg1.bold = True; rg1.font.name = 'Arial'; rg1.font.size = Pt(18); rg1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rg2 = p_g.add_run(f"ARRÊTÉ AU {trimestre.upper()} - EXERCICE {exercice}")
    rg2.bold = True; rg2.font.name = 'Arial'; rg2.font.size = Pt(14); rg2.font.color.rgb = RGBColor(0xD9, 0xE1, 0xF2)

    doc.add_page_break()

    # 1. CADRAGE INITIAL
    _add_heading_arial("1. CADRAGE DU BUDGET INITIAL (CRÉDITS OUVERTS EN LFI)", level=1)
    
    # 1.1 VENTILATION FONCTIONNELLE HIÉRARCHISÉE
    _add_heading_arial("1.1 Optique Opérationnelle : Ventilation Fonctionnelle par Programme & Lignes d'Actions", level=2)
    
    total_rows_fnc = sum(1 + len(p.get("lignes", [])) for p in programmes) + 2
    tbl_fnc = doc.add_table(rows=total_rows_fnc, cols=3)
    tbl_fnc.alignment = WD_TABLE_ALIGNMENT.CENTER
    _appliquer_style_tableau_image(tbl_fnc)

    headers_fnc = ["Programme / Ligne d'Action", "Budget Initial (LFI)", "Poids dans le Programme (%)"]
    for i, h in enumerate(headers_fnc):
        c = tbl_fnc.cell(0, i)
        _set_cell_background(c, HEX_HEADER_BLUE)
        _format_cellule_donnees(c, h, bold=True)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    curr_row = 1
    for p in programmes:
        p_nom = p.get("nom", ""); p_code = p.get("code", "")
        lignes = p.get("lignes", [])
        tot_prog_lfi = sum(float(l.get("lfi", 0) or 0) for l in lignes)

        _set_cell_background(tbl_fnc.cell(curr_row, 0), HEX_SUBHEADER)
        _set_cell_background(tbl_fnc.cell(curr_row, 1), HEX_SUBHEADER)
        _set_cell_background(tbl_fnc.cell(curr_row, 2), HEX_SUBHEADER)

        _format_cellule_donnees(tbl_fnc.cell(curr_row, 0), f"PROGRAMME : {p_code} - {p_nom.upper()}", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        _format_cellule_donnees(tbl_fnc.cell(curr_row, 1), _fmt_fcfa(tot_prog_lfi), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _format_cellule_donnees(tbl_fnc.cell(curr_row, 2), "100,0%", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        curr_row += 1

        for l in lignes:
            lfi_val = float(l.get("lfi", 0) or 0)
            poids_dans_prog = (lfi_val / tot_prog_lfi * 100) if tot_prog_lfi > 0 else 0.0
            _format_cellule_donnees(tbl_fnc.cell(curr_row, 0), f"   └ {l.get('label', '')}", align=WD_ALIGN_PARAGRAPH.LEFT)
            _format_cellule_donnees(tbl_fnc.cell(curr_row, 1), _fmt_fcfa(lfi_val), align=WD_ALIGN_PARAGRAPH.RIGHT)
            _format_cellule_donnees(tbl_fnc.cell(curr_row, 2), _fmt_pct(poids_dans_prog), align=WD_ALIGN_PARAGRAPH.CENTER)
            curr_row += 1

    _format_cellule_donnees(tbl_fnc.cell(curr_row, 0), "TOTAL BUDGET GLOBAL MINISTÉRIEL", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _format_cellule_donnees(tbl_fnc.cell(curr_row, 1), _fmt_fcfa(tot_lfi_global), bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _format_cellule_donnees(tbl_fnc.cell(curr_row, 2), "100,0%", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    _ajouter_source_tableau(doc, "Loi de Finances Initiale (LFI)")
    _ajouter_analyse_paragraphe(doc, "Allocation Opérationnelle", f"La totalité des crédits initiaux ({_fmt_fcfa(tot_lfi_global)}) est mobilisée au profit des 4 programmes opérationnels stratégiques du Ministère.")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.post("/api/rapports/generer")
def generer_rapport(exercice: int = 2026, trimestre: str = "T2"):
    buffer = _generer_rapport_docx(exercice, trimestre)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Rapport_Execution_{trimestre}_{exercice}.docx"'},
    )

@app.post("/api/collecte")
def enregistrer_declaration(payload: SaisieTrimestriellePayload):
    ex = payload.exercice
    if ex in HISTORIQUE_EXERCICES:
        for prog in HISTORIQUE_EXERCICES[ex]:
            if prog["id"] == payload.programme_id:
                for l in prog["lignes"]:
                    match = next((x for x in payload.lignes if x.id == l["id"]), None)
                    if match:
                        l["lfi"], l["ajustement_lfr"], l["engagements"], l["paiements"] = match.lfi, match.ajustement_lfr, match.engagements, match.paiements
    return {"status": "success", "message": f"Saisie {payload.trimestre} - Exercice {ex} enregistrée."}